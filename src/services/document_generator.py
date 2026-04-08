"""
Document generation service - Create Word and PDF documents.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch


class DocumentGenerator:
    """
    Class to handle generating documents (DOCX, PDF) with AI-generated content.
    """

    @staticmethod
    def generate_docx(title, content, filename=None, output_dir="/tmp"):
        """
        Generate a Word document (.docx) with the given title and content.
        
        Args:
            title: Document title
            content: Document content (can include sections separated by ##)
            filename: Optional custom filename (without extension)
            output_dir: Directory to save document
        
        Returns:
            The file path of the generated document
        """
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"document_{timestamp}"
        
        filepath = os.path.join(output_dir, f"{filename}.docx")
        
        doc = Document()
        
        # Add title
        title_heading = doc.add_heading(title, 0)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_heading.runs[0]
        title_run.font.color.rgb = RGBColor(31, 78, 120)
        
        # Add metadata
        metadata_para = doc.add_paragraph()
        metadata_run = metadata_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        metadata_run.italic = True
        metadata_run.font.size = Pt(10)
        
        doc.add_paragraph()  # Add spacing
        
        # Parse content sections
        sections = content.split("##")
        
        for section in sections:
            section = section.strip()
            if not section:
                continue
            
            # Check if this is a section header
            lines = section.split("\n", 1)
            if len(lines) > 1:
                # First line is section title
                section_title = lines[0].strip()
                section_content = lines[1].strip()
                
                # Add section heading
                heading = doc.add_heading(section_title, 1)
                heading_run = heading.runs[0] if heading.runs else None
                if heading_run:
                    heading_run.font.color.rgb = RGBColor(31, 78, 120)
                
                # Add section content
                for line in section_content.split("\n"):
                    line = line.strip()
                    if line:
                        if line.startswith("- ") or line.startswith("• "):
                            # Bullet point
                            doc.add_paragraph(line[2:], style='List Bullet')
                        else:
                            doc.add_paragraph(line)
            else:
                # No section header, just content
                for line in section.split("\n"):
                    line = line.strip()
                    if line:
                        if line.startswith("- ") or line.startswith("• "):
                            doc.add_paragraph(line[2:], style='List Bullet')
                        else:
                            doc.add_paragraph(line)
        
        # Add footer
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "AI-Generated Document"
        
        doc.save(filepath)
        return filepath

    @staticmethod
    def generate_pdf(title, content, filename=None, output_dir="/tmp"):
        """
        Generate a PDF document with the given title and content.
        
        Args:
            title: Document title
            content: Document content (can include sections separated by ##)
            filename: Optional custom filename (without extension)
            output_dir: Directory to save document
        
        Returns:
            The file path of the generated document
        """
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"document_{timestamp}"
        
        filepath = os.path.join(output_dir, f"{filename}.pdf")
        
        # Create PDF document
        doc = SimpleDocTemplate(filepath, pagesize=letter,
                                rightMargin=0.75*inch, leftMargin=0.75*inch,
                                topMargin=0.75*inch, bottomMargin=0.75*inch)
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=RGBColor(31, 78, 120),
            spaceAfter=12,
            alignment=1  # CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=RGBColor(31, 78, 120),
            spaceAfter=10,
            spaceBefore=10
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            spaceAfter=8,
            alignment=4  # JUSTIFY
        )
        
        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=styles['BodyText'],
            fontSize=11,
            leftIndent=20,
            spaceAfter=6
        )
        
        # Build content
        story = []
        
        # Add title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Add metadata
        metadata_text = f"<i>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</i>"
        story.append(Paragraph(metadata_text, styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Parse content sections
        sections = content.split("##")
        
        for section in sections:
            section = section.strip()
            if not section:
                continue
            
            lines = section.split("\n", 1)
            if len(lines) > 1:
                # First line is section title
                section_title = lines[0].strip()
                section_content = lines[1].strip()
                
                # Add section heading
                story.append(Paragraph(section_title, heading_style))
                
                # Add section content
                for line in section_content.split("\n"):
                    line = line.strip()
                    if line:
                        if line.startswith("- ") or line.startswith("• "):
                            # Bullet point
                            story.append(Paragraph(f"• {line[2:]}", bullet_style))
                        else:
                            story.append(Paragraph(line, body_style))
            else:
                # No section header, just content
                for line in section.split("\n"):
                    line = line.strip()
                    if line:
                        if line.startswith("- ") or line.startswith("• "):
                            story.append(Paragraph(f"• {line[2:]}", bullet_style))
                        else:
                            story.append(Paragraph(line, body_style))
        
        # Build PDF
        doc.build(story)
        return filepath
